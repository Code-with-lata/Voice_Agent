import streamlit as st
import os
import re
import base64
import tempfile
import time
from faster_whisper import WhisperModel
from groq import Groq
from gtts import gTTS
from docx import Document
import PyPDF2
from dotenv import load_dotenv
from docx.shared import Pt
from audio_recorder_streamlit import audio_recorder


load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
LLM_MODEL = "llama-3.1-8b-instant"

# ================== SMART PROMPT ==================
SYSTEM_PROMPT = (
    "You are a friendly and professional AI Technical Interviewer.\n\n"
    
    "STRICT FLOW:\n"
    
    "1. FIRST MESSAGE (VERY IMPORTANT):\n"
    "- Start with a warm human-like greeting.\n"
    "- Briefly introduce the interview process.\n"
    "- Ask if the candidate is ready.\n"
    "- DO NOT ask any technical question yet.\n\n"

    "2. AFTER USER SAYS YES:\n"
    "- Ask the candidate to briefly introduce themselves.\n"
    "- Do NOT start technical questions yet.\n\n"
    
    "3. AFTER USER RESPONDS:\n"
    "- Give a short acknowledgment (2-5 words).\n"
    "- If user confirms (yes/ready), start interview with first question.\n"
    
    "4. DURING INTERVIEW:\n"
    "- Evaluate answers internally.\n"
    "- Give short feedback (2-5 words).\n"
    "- Then ask next question OR follow-up.\n\n"
    
    "5. STYLE:\n"
    "- Keep responses under 20 words.\n"
    "- Be natural and human-like.\n"
    "- Avoid repetition.\n"
)

# ================== FILE PARSING ==================
def extract_questions_from_file(uploaded_file):
    questions = []
    file_type = uploaded_file.name.split('.')[-1].lower()

    if file_type == 'docx':
        doc = Document(uploaded_file)
        questions = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    elif file_type == 'pdf':
        reader = PyPDF2.PdfReader(uploaded_file)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                questions.extend([line.strip() for line in text.split('\n') if line.strip()])
    
    return questions


# ================== VOICE FUNCTIONS ==================
@st.cache_resource
def load_agent_engines():
    stt = WhisperModel("base", device="cpu", compute_type="int8")
    client = Groq(api_key=GROQ_API_KEY)
    return stt, client


def generate_greeting(client, history):
    response = client.chat.completions.create(
        model=LLM_MODEL,
        messages=history + [{"role": "user", "content": "Start the interview."}],
        temperature=0.7
    )
    return response.choices[0].message.content


def get_ai_decision(client, user_text, next_q, history):
    """AI Brain: Follow-up ya Next Question ka faisla"""
    prompt = f"""
                Candidate said: "{user_text}". If follow-up is needed, ask it. Otherwise, ask the 
                next question: "{next_q}"

                INSTRUCTION: 
                - Briefly validate if the answer is correct (DON'T repeat it).
                - If the answer is good, ask the Next Scheduled Question.
                - If the answer is technically flawed or too short, ask a specific follow-up about that topic.

                
                """
    response = client.chat.completions.create(
        model=LLM_MODEL,
        messages=history + [{"role": "system", "content": prompt}],
        temperature=0.7
    )
    return response.choices[0].message.content


def ai_voice_output(text):
    """AI voice output with auto-play"""
    if not text:
        return
        
    try:
        tts = gTTS(text=text, lang='en')
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as fp:
            tts.save(fp.name)
            with open(fp.name, "rb") as f:
                data = f.read()
                b64 = base64.b64encode(data).decode()
            
            unique_id = f"audio_{int(time.time() * 1000)}"
            
            audio_html = f"""
                <audio id="{unique_id}" autoplay="true" style="display:none;">
                    <source src="data:audio/mp3;base64,{b64}" type="audio/mp3">
                </audio>
                <script>
                    var audio = document.getElementById('{unique_id}');
                    audio.onended = function() {{
                        window.parent.postMessage({{type: 'streamlit:setComponentValue', value: 'finished'}}, '*');
                    }};
                    audio.play();
                </script>
            """
            st.components.v1.html(audio_html, height=0)
        os.remove(fp.name)
    except Exception as e:
        st.error(f"Voice Error: {e}")



def transcribe_audio(stt_model, audio_bytes):
    """Audio bytes ko text me convert karo"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
            tmp.write(audio_bytes)
            tmp.flush()
            
            segments, _ = stt_model.transcribe(tmp.name)
            text = " ".join([s.text for s in segments])
            
        os.remove(tmp.name)
        return text.strip()
    except Exception as e:
        return ""


# ================== PDF REPORT ==================

def generate_report():
    doc = Document()
    doc.add_heading('AI Interview Performance Report', 0)
    
    # Overall Score Calculation (e.g., 7 / 15)
    total_obtained = sum(st.session_state.scores)
    max_possible = len(st.session_state.scores) * 5
    
    summary = doc.add_paragraph()
    run = summary.add_run(f"TOTAL INTERVIEW SCORE: {total_obtained} / {max_possible}")
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph("_" * 40)
    doc.add_heading('Detailed Evaluation:', level=1)

    if not st.session_state.answers:
        doc.add_paragraph("No questions were answered.")
    else:
        for i, item in enumerate(st.session_state.answers):
            # Question Heading
            q = doc.add_paragraph()
            q.add_run(f"Question {i+1}: {item['question']}").bold = True
            
            # User's Answer
            doc.add_paragraph(f"Your Answer: {item['answer']}")
            
            # Simple Score Display (Correctness, Communication, Confidence )
            s = doc.add_paragraph()
            s.add_run(f"Final Score: {item['final_score']} / 5").italic = True
            
            # Divider line
            doc.add_paragraph("-" * 30)

    file_path = "interview_report.docx"
    doc.save(file_path)
    return file_path


# ================== SESSION MANAGEMENT ==================
def init_session():
    defaults = {
        "chat_history": [{"role": "system", "content": SYSTEM_PROMPT}],
        "q_bank": [],
        "q_index": 0,
        "mic_counter": 0,
        "is_started": False,
        "awaiting_intro": False,
        "answers": [],
        "scores": [],
        "report_ready": False,
        "report_file": None,
        "force_end": False,
        "followup_count": 0,
        "pending_voice": None,
        "last_audio_id": None
    }
    
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


# ================== UI SCREENS ==================
def render_upload_screen(groq_client):
    st.markdown("""
    <div style="text-align: center; padding: 30px;">
        <h2>🎯 AI-Powered Voice Interview</h2>
        <p style="font-size: 18px; color: #666;">Real-time conversation with AI Interviewer</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.info("📋 Upload your interview questions (PDF or DOCX)")
    file = st.file_uploader("", type=['pdf', 'docx'], label_visibility="collapsed")
    
    if file:
        if st.button("🚀 Start Interview", use_container_width=True, type="primary"):
            questions = extract_questions_from_file(file)
            
            if questions:
                greeting = generate_greeting(groq_client, st.session_state.chat_history)
                st.session_state.q_bank = questions
                st.session_state.is_started = True
                st.session_state.chat_history.append({"role": "assistant", "content": greeting})
                st.session_state.pending_voice = greeting
                st.session_state.first_question = questions[0]
                st.rerun()
            else:
                st.error("❌ No questions found in file!")



def process_user_audio(audio_bytes, stt_model, groq_client):
    """Process user's audio response"""
    
    # Create unique ID for this audio
    audio_id = hash(audio_bytes)
    
    # Skip if already processed
    if audio_id == st.session_state.last_audio_id:
        return
    
    st.session_state.last_audio_id = audio_id
    
    with st.spinner("🎯 Processing your response..."):
        user_text = transcribe_audio(stt_model, audio_bytes)
        
        if not user_text or len(user_text) < 3:
            st.warning("⚠️ Could not hear clearly. Please try again.")
            st.session_state.last_audio_id = None
            return
        
        # Add to chat
        st.session_state.chat_history.append({"role": "user", "content": user_text})
        

        
        # ===== GREETING/INTRO FLOW =====
        if "first_question" in st.session_state and not st.session_state.awaiting_intro:
            if any(word in user_text.lower() for word in ["yes", "ready", "start", "ok", "sure", "let's"]):
                reply = "Perfect! Please introduce yourself briefly."
                st.session_state.awaiting_intro = True
            else:
                reply = "No worries. Take your time. Ready when you are!"
        
        elif st.session_state.get("awaiting_intro"):
            reply = f"Great to meet you! Let's begin. {st.session_state.first_question}"
            st.session_state.q_index = 1
            st.session_state.awaiting_intro = False
            del st.session_state.first_question
        
        # ===== INTERVIEW FLOW =====
        else:
            idx = st.session_state.q_index
            q_bank = st.session_state.q_bank
            
            # Check if interview complete
            if idx >= len(q_bank) and len(st.session_state.answers) > 0:
                reply = "Excellent work! Interview complete. Generating your report..."
                pdf = generate_report()
                st.session_state.report_ready = True
                st.session_state.report_file = pdf
            
            else:
                current_q = q_bank[idx-1] if idx > 0 else q_bank[0]
                next_q = q_bank[idx] if idx < len(q_bank) else "End"
                
                # Score the answer based on three aspects
                score_prompt = f"""
                    Question: {current_q}
                    Answer: {user_text}
                
                    Evaluate the answer on a scale of 0-5 for each of these categories:
                    1. Correctness: Is the technical information accurate?
                    2. Communication: Is the explanation clear and well-structured?
                    3. Confidence: Does the tone/content reflect certainty?
                
                    Return the result ONLY in this exact format:
                    Correctness: [score], Communication: [score], Confidence: [score]
                """
                
                try:
                    score_res = groq_client.chat.completions.create(
                        model=LLM_MODEL,
                        messages=[{"role": "user", "content": score_prompt}],
                        temperature=0.3
                    )
                
                    res_text = score_res.choices[0].message.content.strip()
    
                    scores = re.findall(r'\d+', res_text)
                    scores = [int(s) for s in scores]
                    
                    # Teeno aspects ka average (Final Question Score)
                    question_final_score = round(sum(scores) / 3, 1) if scores else 0
                
                
                    # Save only valid interview answers
                    if user_text.strip():
                        st.session_state.answers.append({
                            "question": current_q,
                            "answer": user_text.strip(),
                            #  "score": score
                            "final_score": question_final_score  
                        })
                        
                        st.session_state.scores.append(question_final_score)
                
                except Exception as e:
                    st.error(f"Scoring Error: {e}")
                    
                # Get AI response
                ai_reply = get_ai_decision(groq_client, user_text, next_q, st.session_state.chat_history)
                reply = ai_reply
                
                # Check if moved to next question
                if next_q.lower() in reply.lower() or "next" in reply.lower():
                    st.session_state.q_index += 1
                    st.session_state.followup_count = 0
                
                # Detect follow-up
                elif any(word in reply.lower() for word in ["why", "how", "example", "explain", "elaborate", "can you"]):
                    st.session_state.followup_count += 1
                
                # Max 2 follow-ups, then move on
                if st.session_state.followup_count >= 2:
                    reply = f"Understood. Let's move on. {next_q}"
                    st.session_state.q_index += 1
                    st.session_state.followup_count = 0
        
        # Add AI response
        st.session_state.chat_history.append({"role": "assistant", "content": reply})
        st.session_state.pending_voice = reply
        # st.session_state.mic_counter += 1  # 🔥 important
        
        st.rerun()


# ================== MAIN APP ==================
def main():
    st.set_page_config(
        page_title="🎙️ AI-Powered Voice-Based Interview Assistant",
        page_icon="🎤",
        layout="centered"
    )
    
    # Custom CSS
    st.markdown("""
    <style>
        .main {
            background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        }
        .stButton > button {
            border-radius: 10px;
            font-weight: 600;
        }
    </style>
    """, unsafe_allow_html=True)
    
    st.title("🎙️ AI Voice Interview")
    
    # Initialize
    init_session()
    stt_model, groq_client = load_agent_engines()
    
    # ===== UPLOAD SCREEN =====
    if not st.session_state.is_started:
        render_upload_screen(groq_client)
        return
    
    # ===== INTERVIEW SCREEN =====
    
    # Play AI voice if pending
    if st.session_state.pending_voice:
        ai_voice_output(st.session_state.pending_voice)

        wait_time = (len(st.session_state.pending_voice) / 10) + 2

        time.sleep(wait_time)

        st.session_state.pending_voice = None

        st.session_state.mic_counter += 1 

        st.rerun()
    
    # Control panel
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        progress = min(st.session_state.q_index, len(st.session_state.q_bank))
        total = len(st.session_state.q_bank)
        st.metric("Progress", f"{progress}/{total}")
    
    with col3:
        if st.button("🛑 End Interview", type="secondary"):
            st.session_state.force_end = True
            st.rerun()
    
    # Handle force end
    if st.session_state.force_end:
        
        st.session_state.chat_history.append({
            "role": "assistant",
            "content": "Interview ended early. Generating report..."
        })
        pdf = generate_report()
        st.session_state.report_ready = True
        st.session_state.report_file = pdf
        st.session_state.pending_voice = "Interview ended. Your report is ready."
        st.session_state.force_end = False
        st.rerun()
    
    st.markdown("---")
    
    
    # 3. AUTOMATIC MIC (AI ke chup hone ke baad)
    if not st.session_state.report_ready: 
        if st.session_state.pending_voice is None:
            st.write("### 🎤 AI is listening... (Speak now)")
            
            audio_bytes = audio_recorder(
                text="Listening...",
                recording_color="#e74c3c",
                neutral_color="#3498db",
                icon_name="microphone",
                icon_size="2x",
                pause_threshold=2.5, 
                sample_rate=16000,
                auto_start=True,
                key=f"mic_{st.session_state.mic_counter}" # Unique key for each turn
            )

        if audio_bytes:
            process_user_audio(audio_bytes, stt_model, groq_client)

    
    # ===== DOWNLOAD REPORT =====
    if st.session_state.report_ready and st.session_state.report_file:
        st.success("✅ Interview Completed!")
        
        total = sum(st.session_state.scores)
        maximum = len(st.session_state.scores) * 5
        percentage = (total / maximum * 100) if maximum > 0 else 0
        
        st.metric("Final Score", f"{total}/{maximum}", f"{percentage:.1f}%")
        
        with open(st.session_state.report_file, "rb") as f:
            st.download_button(
                label="📄 Download Report ",
                data=f,
                file_name="interview_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )


if __name__ == "__main__":
    main()
