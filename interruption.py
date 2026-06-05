import streamlit as st
import os
import re
import base64
import tempfile
import time
from faster_whisper import WhisperModel
from groq import Groq
import asyncio  
import edge_tts
# from gtts import gTTS
from docx import Document
import PyPDF2
from dotenv import load_dotenv
from docx.shared import Pt
from audio_recorder_streamlit import audio_recorder


load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
LLM_MODEL = "llama-3.1-8b-instant"

# ================== SMART PROMPT ==================
SYSTEM_PROMPT = (
    "You are a friendly and professional AI Technical Interviewer.\n\n"
    
    "STRICT FLOW:\n"
    
    "1. FIRST MESSAGE (VERY IMPORTANT):\n"
    "- Start with a warm human-like greeting.\n"
    "- Briefly introduce the interview process.\n"
    "- Ask if the candidate is ready.\n"
    "- DO NOT ask any technical question yet.\n\n"

    "2. AFTER USER SAYS YES:\n"
    "- Ask the candidate to briefly introduce themselves.\n"
    "- Do NOT start technical questions yet.\n\n"
    
    "3. AFTER USER RESPONDS:\n"
    "- Give a short acknowledgment (2-5 words).\n"
    "- If user confirms (yes/ready), start interview with first question.\n"
    
    "4. DURING INTERVIEW:\n"
    "- Evaluate answers internally.\n"
    "- Give short feedback (2-5 words).\n"
    "- Then ask next question OR follow-up.\n\n"
    
    "5. STYLE:\n"
    "- Keep responses under 20 words.\n"
    "- Be natural and human-like.\n"
    "- Avoid repetition.\n"
)

# ================== FILE PARSING ==================
def extract_questions_from_file(uploaded_file):
    questions = []
    file_type = uploaded_file.name.split('.')[-1].lower()

    if file_type == 'docx':
        doc = Document(uploaded_file)
        questions = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    elif file_type == 'pdf':
        reader = PyPDF2.PdfReader(uploaded_file)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                questions.extend([line.strip() for line in text.split('\n') if line.strip()])
    
    return questions


# ================== VOICE FUNCTIONS ==================
@st.cache_resource
def load_agent_engines():
    stt = WhisperModel("medium", device="cpu", compute_type="int8")
    
    client = Groq(api_key=GROQ_API_KEY)
    return stt, client


def generate_greeting(client, history):
    response = client.chat.completions.create(
        model=LLM_MODEL,
        messages=history + [{"role": "user", "content": "Start the interview."}],
        temperature=0.7
    )
    return response.choices[0].message.content


def get_ai_decision(client, user_text, next_q, history):
    """AI Brain: Follow-up ya Next Question ka faisla"""
    prompt = f"""
                Candidate said: "{user_text}". If follow-up is needed, ask it. Otherwise, ask the 
                next question: "{next_q}"

                INSTRUCTION: 
                1. If answer is correct:
                - briefly appreciate
                - ask next interview question
                
                2. If answer is partially correct:
                - ask a follow-up question
                
                3. If candidate says:
                - "don't know"
                - "no idea"
                - silence
                THEN:
                - encourage briefly
                - move to next question
                
                IMPORTANT:
                - NEVER explain the answer yourself
                - NEVER teach concepts
                - NEVER give definitions
                - Keep response under 25 words

                
                """
    response = client.chat.completions.create(
        model=LLM_MODEL,
        messages=history + [{"role": "system", "content": prompt}],
        temperature=0.7
    )
    return response.choices[0].message.content


# ================== SPEECH CLASSIFIER ==================

def classify_user_input(client, text, current_question):

    prompt = f"""
You are a STRICT real-time AI interview speech classifier.

CURRENT INTERVIEW QUESTION:
"{current_question}"

USER SPEECH:
"{text}"

Classify into EXACTLY ONE category:

1. TECHNICAL_ANSWER
2. INTERRUPTION
3. BACKGROUND_TALK
4. NEXT_QUESTION

RULES:

TECHNICAL_ANSWER:
- Interview-related technical answer
- Concepts/coding/examples
- Explanations related to the interview question

INTERRUPTION:
- Asking to repeat
- Asking to slow down

Examples:
- repeat please
- sorry
- come again
- what
- again
- can you repeat
- speak slowly

BACKGROUND_TALK:
- Talking to someone else
- Greetings
- Names
- Phone calls
- Side conversations

Examples:
- hello rahul
- mummy ek minute
- phone aa raha hai
- hold on
- hello sir
- someone is helping me
- wait bro
- haan bata
- kya answer h
- google kar
- read this

NEXT_QUESTION:
- User wants to skip current question
- User wants another question
- User wants to move ahead

Examples:
- next question
- move to next
- skip this
- let's continue
- move ahead
- go to next question
- ask another question

IMPORTANT:
- If user wants another question, ALWAYS return NEXT_QUESTION
- If unsure, choose BACKGROUND_TALK
- Return ONLY category name
- No explanation

Valid outputs:
TECHNICAL_ANSWER
INTERRUPTION
BACKGROUND_TALK
NEXT_QUESTION
"""

    try:

        response = client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": "You are a strict speech classifier."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0
        )

        result = response.choices[0].message.content.strip()

        valid = [
            "TECHNICAL_ANSWER",
            "INTERRUPTION",
            "BACKGROUND_TALK",
            "NEXT_QUESTION"
        ]

        if result not in valid:
            return "BACKGROUND_TALK"

        return result

    except:
        return "BACKGROUND_TALK"


async def generate_edge_voice(text, output_path):
    """Edge-TTS se high quality audio generate karne wala function"""
    # Aap 'en-IN-PrabhatNeural' (Male) ya 'en-IN-NeerjaNeural' (Female) use kar sakte hain
    communicate = edge_tts.Communicate(text, "en-IN-PrabhatNeural")
    await communicate.save(output_path)


def ai_voice_output(text):
    if not text:
        return
        
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as fp:
            temp_name = fp.name
            
        asyncio.run(generate_edge_voice(text, temp_name))
        
        with open(temp_name, "rb") as f:
            data = f.read()
            b64 = base64.b64encode(data).decode()
        
        unique_id = f"audio_{int(time.time() * 1000)}"
        
        audio_html = f"""
            <div id="status_{unique_id}" style="color: #764ba2; font-size: 14px; font-weight: bold;">
                🔊 AI is speaking... (Try interrupting me)
            </div>
            <audio id="{unique_id}" autoplay>
                <source src="data:audio/mp3;base64,{b64}" type="audio/mp3">
            </audio>
            
            <script>
                (function() {{
                    const audio = document.getElementById('{unique_id}');
                    const status = document.getElementById('status_{unique_id}');
                    
                    let playbackStarted = false;
                    let interruptionDetected = false;
                    let echoGuardTime = Date.now() + 600; // 600ms ka buffer

                    audio.play().then(() => {{
                        playbackStarted = true;
                    }}).catch(e => console.log(e));

                    async function detectVoice() {{
                        try {{
                            const stream = await navigator.mediaDevices.getUserMedia({{ audio: true }});
                            const audioContext = new AudioContext();
                            const source = audioContext.createMediaStreamSource(stream);
                            const analyser = audioContext.createAnalyser();
                            analyser.fftSize = 256;
                            source.connect(analyser);

                            const bufferLength = analyser.frequencyBinCount;
                            const dataArray = new Uint8Array(bufferLength);

                            function checkVolume() {{
                                analyser.getByteFrequencyData(dataArray);
                                let values = 0;
                                let highPeaks = 0;
                                for (let i = 0; i < bufferLength; i++) {{
                                    values += dataArray[i];

                                    if (dataArray[i] > 180) {{
                                        highPeaks++;
                                    }}
                                }}
                                let average = values / bufferLength;

                                // MULTIPLE VOICE DETECTION
                                if (highPeaks > 25) {{
        
                                    console.log("Possible multiple voices detected");
        
                                    localStorage.setItem("multiple_voice", "true");
        
                                    status.innerHTML = "⚠️ Multiple voices detected";
                                    status.style.color = "red";
                                }}

                                // ECHO GUARD LOGIC:
                                
                                if (playbackStarted && Date.now() > echoGuardTime && average > 55 && !interruptionDetected ) {{ 
                                    interruptionDetected = true;
                                    console.log("True User Interruption Detected!");
                                    audio.pause();
                                    status.innerHTML = "🎤 Listening...";
                                    status.style.color = "red";
                                }}

                                // TAB SWITCH DETECTION
                                document.addEventListener("visibilitychange", () => {{
        
                                    if (document.hidden) {{
                                        localStorage.setItem("tab_switch", "true");
                                    }}
                                }});
                                        
                                if (!audio.paused && !audio.ended) {{
                                    requestAnimationFrame(checkVolume);
                                }} else if (audio.ended) {{
                                    status.innerHTML = "✅ Finished speaking";
                                }}
                            }}
                            checkVolume();
                        }} catch (err) {{
                            console.error("Mic error:", err);
                        }}
                    }}
                    
                    detectVoice();
                }})();
            </script>
        """
        st.components.v1.html(audio_html, height=60)
        
        if os.path.exists(temp_name):
            try:
                os.remove(temp_name)
            except:
                pass
            
    except Exception as e:
        st.error(f"Edge-Voice Error: {e}")

def transcribe_audio(stt_model, audio_bytes):
    """Audio bytes ko text me convert karo"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
            tmp.write(audio_bytes)
            tmp.flush()
            
            segments, info = stt_model.transcribe(
                tmp.name,
                beam_size=5,
                best_of=5,
                language="en",
                temperature=0,
                vad_filter=True,
                vad_parameters=dict(min_silence_duration_ms=500),
            )
            text = " ".join([s.text for s in segments])
            
        os.remove(tmp.name)
        return text.strip()
    except Exception as e:
        return ""
#
# -------------------- Report Generation Logic --------------------
def generate_report():
    doc = Document()
    doc.add_heading('AI Interview: Final Performance Report', 0)

    client = Groq(api_key=GROQ_API_KEY)
    
    # ------------------  CHAT CONVERSATION ----------------------
    doc.add_heading('Interview Conversation:', level=1)
    transcript_text = ""
    for message in st.session_state.chat_history:
        
        if message['role'] == "system": 
            continue
        p = doc.add_paragraph()
        role_label = "Interviewer (Bot): " if message['role'] == "assistant" else "Candidate (User): "
        transcript_text += f"{role_label}{message['content']}\n"
        run = p.add_run(f"{role_label}{message['content']}")
    
        if message['role'] == "assistant": 
            run.bold = True
        else:
            run.bold = False

    doc.add_page_break() 


    # ---  TECHNICAL SCORING (Performing Scoring Here) ---
    doc.add_heading('Technical Performance Evaluation', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Question'; hdr_cells[1].text = 'Evaluation'; hdr_cells[2].text = 'Score'

    calculated_scores = []

    for q in st.session_state.q_bank:
        item = next((a for a in st.session_state.answers if a["question"] == q), None)
        
        if item:
            main_answer = item.get("main_answer", "")
            followups = item.get("followups", [])
        
            answer_text = (
                main_answer +
                "\nFollow-up Responses:\n" +
                "\n".join(followups)
            )
        else:
            answer_text = ""

        # ❌ Invalid answer → direct 0
        word_count = len(answer_text.split())
        if answer_text.strip() in ["...", ".", ""]: # or len(answer_text.split()) < 3:
            q_score = 0.0
            
        else:
            score_prompt = f"""
                You are a STRICT senior technical interviewer.
                
                QUESTION:
                {q}
                
                CANDIDATE ANSWER:
                {answer_text}
                
                STRICT SCORING RULES:
                
                - Give HIGH scores ONLY if technically deep and accurate.
                - Surface-level answers MUST score low.
                - Wrong technical facts = 0.
                - Short answers (<15 words) cannot score above 2.
                - If answer lacks examples, reasoning, implementation details,
                  edge cases, or practical understanding, reduce Depth score.
                - Avoid generosity.
                
                Evaluate from 0-5 in these categories:
                
                1. Technical Correctness
                2. Depth & Understanding
                3. Technical Terminology
                4. Communication Clarity
                
                Also provide a 1-line professional evaluation.
                
                Return EXACTLY in this format:
                
                Correctness: X
                Depth: X
                Keywords: X
                Communication: X
                Evaluation: your feedback
            """
    
            try:
                res = client.chat.completions.create(
                    model=LLM_MODEL,
                    messages=[{"role": "system", "content": "You are a strict technical interviewer. Return only numeric scores."},
                              {"role": "user", "content": score_prompt}],
                    temperature=0.1
                )
                res_text = res.choices[0].message.content
                scores = re.findall(
                    r'Correctness:\s*(\d).*?Depth:\s*(\d).*?Keywords:\s*(\d).*?Communication:\s*(\d)',
                    res_text,
                    re.IGNORECASE | re.DOTALL
                )
                evaluation_match = re.search(
                    r'Evaluation:\s*(.*)',
                    res_text,
                    re.IGNORECASE | re.DOTALL
                )
                
                if scores:
                        nums = list(map(int, scores[0]))
                        q_score = round(sum(nums) / 4, 1)
                else:
                    q_score = 0.0

                evaluation_text = (
                    evaluation_match.group(1).strip()
                    if evaluation_match
                    else "Evaluation completed."
                ) 
            except Exception as e:
                print(f"Error scoring: {e}")
                q_score = 0.0
                evaluation_text = "Evaluation failed."
    
            calculated_scores.append(q_score)
    
            row_cells = table.add_row().cells
            row_cells[0].text = q
            row_cells[1].text = evaluation_text #"Technical assessment completed"
            row_cells[2].text = f"{q_score} / 5"

    st.session_state.scores = calculated_scores    

    # ---  CONVERSATION SKILLS SCORING ---
    doc.add_heading('2. Communication & Soft Skills Analysis', level=1)
    
    conv_prompt = f"""
        Analyze the following interview transcript.
        
        TRANSCRIPT:
        {transcript_text}
        
        Evaluate the candidate on:
        
        1. Communication Skills
        2. Professionalism
        3. Confidence
        
        SCORING RULES:
        - Be strict and realistic
        - Poor communication should score low
        - Hesitation, silence, or vague speaking should reduce score
        
        Return EXACTLY in this format:
        
        Communication: X/10
        Professionalism: X/10
        Confidence: X/10
        Feedback: your detailed analysis
    """
    
    try:
        res = client.chat.completions.create(
            model=LLM_MODEL,
            messages=[{"role": "system", "content": "You are a senior HR manager."},
                      {"role": "user", "content": conv_prompt}],
            temperature=0.3
        )
        conv_analysis = res.choices[0].message.content
        doc.add_paragraph(conv_analysis)
        # Extracting a numeric score for calculation (assuming format "Score: 8")
        comm = re.search(r'Communication.*?(\d+)/10', conv_analysis, re.IGNORECASE)
        prof = re.search(r'Professionalism.*?(\d+)/10', conv_analysis, re.IGNORECASE)
        conf = re.search(r'Confidence.*?(\d+)/10', conv_analysis, re.IGNORECASE)
    
        scores = []
    
        if comm:
            scores.append(int(comm.group(1)))
        if prof:
            scores.append(int(prof.group(1)))
        if conf:
            scores.append(int(conf.group(1)))
    
        # Final average score
        if scores:
            conv_score = sum(scores) / len(scores)  
        else:
            conv_score = 5 
    except Exception as e:
        conv_score = 5 
        doc.add_paragraph("Soft skills evaluation: Professional and clear communication observed.")

    
    # ================= INTEGRITY ANALYSIS =================

    doc.add_heading('Interview Integrity Analysis', level=1)
    
    if st.session_state.cheating_flags:
    
        for flag in st.session_state.cheating_flags:
            doc.add_paragraph(f"⚠️ {flag}")
    
        doc.add_paragraph(
            f"Suspicion Score: {st.session_state.suspicion_score}"
        )
    
    else:
        doc.add_paragraph("No suspicious activity detected.")

    # --- COMPLETE SUMMARY & FINAL TOTAL SCORE ---
    doc.add_heading(' Executive Summary:', level=1)
    total_obtained = sum(calculated_scores)
    total_questions = len(st.session_state.q_bank)
    max_possible = total_questions * 5  #len(calculated_scores) * 5

    percentage = (total_obtained / max_possible * 100) if max_possible > 0 else 0
    conv_percent = (conv_score / 10 * 100)
    final_weighted_score = (percentage * 0.7) + (conv_percent * 0.3)

    # # Summary Logic based on score
    # if percentage >= 80:
    #     performance_level = "EXCELLENT"
    #     remarks = "The candidate shows strong technical command and clear communication. Highly recommended."
    # elif percentage >= 60:
    #     performance_level = "GOOD"
    #     remarks = "The candidate has a solid foundation but could improve on technical depth in certain areas."
    # else:
    #     performance_level = "NEEDS IMPROVEMENT"
    #     remarks = "The candidate struggled with core concepts. Further training or review is suggested."
    # doc.add_paragraph("\n") # Space

    summary_p = doc.add_paragraph()
    summary_p.add_run(f"FINAL INTERVIEW SCORE: {total_obtained:.1f} / {max_possible}").bold = True
    summary_p.add_run(f"\nPERCENTAGE: {percentage:.1f}%").bold = True
    summary_p.add_run(f"\nCommunication Score: {conv_percent:.1f}%").bold = True
    summary_p.add_run(f"\n\nOVERALL INTERVIEW RATING: {final_weighted_score:.1f}%").bold = True
    # summary_p.add_run(f"\nPERFORMANCE STATUS: {performance_level}").bold = True
    
    status = "SELECTED" if final_weighted_score >= 75 else "RE-EVALUATE" if final_weighted_score >= 50 else "REJECTED"
    summary_p.add_run(f"\nRESULT: {status}").bold = True

    # Feedback basis on percentage
    # doc.add_heading('Interviewer Remarks:', level=2)
    # doc.add_paragraph(remarks)

    file_path = "interview_report.docx"
    doc.save(file_path)
    return file_path


# ================== SESSION MANAGEMENT ==================
def init_session():
    defaults = {
        "chat_history": [{"role": "system", "content": SYSTEM_PROMPT}],
        "q_bank": [],
        "q_index": 0,
        "mic_counter": 0,
        "is_started": False,
        "awaiting_intro": False,
        "answers": [],
        "scores": [],
        "report_ready": False,
        "report_file": None,
        "force_end": False,
        "followup_count": 0,
        "pending_voice": None,
        "last_audio_id": None,
        "pause_count": 0,  
        "is_paused_state": False,
        "current_question": None,
        "cheating_flags": [],
        "suspicion_score": 0,
        "tab_switch_count": 0,
        "multiple_voice_count": 0,
        "is_ai_speaking": False,
        "last_response_time": time.time()
        
    }
    
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


# ================== UI SCREENS ==================
def render_upload_screen(groq_client):
    st.markdown("""
    <div style="text-align: center; padding: 30px;">
        <h2>🎯 AI-Powered Voice Interview</h2>
        <p style="font-size: 18px; color: #666;">Real-time conversation with AI Interviewer</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.info("📋 Upload your interview questions (PDF or DOCX)")
    file = st.file_uploader("", type=['pdf', 'docx'], label_visibility="collapsed")
    
    if file:
        if st.button("🚀 Start Interview", use_container_width=True, type="primary"):
            questions = extract_questions_from_file(file)
            
            if questions:
                greeting = generate_greeting(groq_client, st.session_state.chat_history)
                st.session_state.q_bank = questions
                st.session_state.is_started = True
                st.session_state.chat_history.append({"role": "assistant", "content": greeting})
                st.session_state.pending_voice = greeting
                st.session_state.first_question = questions[0]
                st.rerun()
            else:
                st.error("❌ No questions found in file!")



def process_user_audio(audio_bytes, stt_model, groq_client):
    
    audio_id = hash(audio_bytes)
    if audio_id == st.session_state.last_audio_id:
        return
    st.session_state.last_audio_id = audio_id

    with st.spinner("🎯 Processing your response..."):

        user_text = transcribe_audio(stt_model, audio_bytes)

        # AI Speaking & Interruption Logic
        if st.session_state.get("is_ai_speaking") and time.time() < st.session_state.get("ai_end_time", 0):
    
            if not user_text or len(user_text.strip().split()) < 3:
                st.session_state.mic_counter += 1
                st.rerun()
                return
            else:
                st.session_state.is_ai_speaking = False
                st.session_state.ai_end_time = 0

        if not user_text or len(user_text.strip()) == 0:
            
            user_text = "..."
        
        response_delay = time.time() - st.session_state.last_response_time
        word_count = len(user_text.split())
        
        if response_delay > 10 and word_count > 80:
            st.session_state.suspicion_score += 1
            st.session_state.cheating_flags.append(
                "Long silence followed by unusually detailed answer"
            )

        st.session_state.last_response_time = time.time()

        # SAVE USER MESSAGE 
        st.session_state.chat_history.append({
            "role": "user",
            "content": user_text
        })

        idx = st.session_state.q_index
        q_bank = st.session_state.q_bank
        current_q = st.session_state.current_question

        # PAUSE HANDLING 

        if user_text.strip() in ["", "..."]:

            st.session_state.pause_count += 1

            if st.session_state.pause_count == 1:
                reply = "I noticed a silence. Do you need a moment to think, or should we move to the next question?"

            elif st.session_state.pause_count == 2:
                reply = "I again noticed a silence. No problem, take your time to collect your thoughts. I'm still listening."

            else:
                next_q = q_bank[idx] if idx < len(q_bank) else "Interview completed."
                reply = f"I'm sorry, but we have already spent quite some time here. To ensure we cover everything, I'm moving to the next question. {next_q}"

                st.session_state.pause_count = 0

                if idx < len(q_bank):
                    st.session_state.q_index += 1
                    st.session_state.current_question = next_q

            st.session_state.chat_history.append({
                "role": "assistant",
                "content": reply
            })

            st.session_state.pending_voice = reply
            st.rerun()

        st.session_state.pause_count = 0

        # INTRO FLOW 

        positive_reply = any(word in user_text.lower() for word in [
            "yes", "ready", "ok", "sure", "start"
        ])
        
        negative_reply = any(word in user_text.lower() for word in [
            "no", "not now", "later"
        ]) or "not now" in user_text.lower()

        if (
            "first_question" in st.session_state
            and not st.session_state.awaiting_intro
            and not st.session_state.get("intro_completed", False)
            and st.session_state.current_question is None
        ):

            if positive_reply:
                reply = "Perfect! Please introduce yourself briefly."
                st.session_state.awaiting_intro = True
            elif negative_reply:
                reply = "No worries. Take your time. Tell me when you're ready to start!"    
            else:
                reply = "Please say yes when you're ready."

            st.session_state.chat_history.append({
                "role": "assistant",
                "content": reply
            })

            st.session_state.pending_voice = reply
            st.rerun()

        if st.session_state.get("awaiting_intro") and len(user_text.split()) > 3:

            q = st.session_state.first_question
            st.session_state.current_question = q
            st.session_state.q_index = 1
            st.session_state.awaiting_intro = False
            st.session_state.intro_completed = True

            reply = f"Great to meet you! Let's jump into the first question. {q}"

            st.session_state.chat_history.append({
                "role": "assistant",
                "content": reply
            })

            st.session_state.pending_voice = reply
            st.rerun()

        # 🔥 SINGLE LLM CLASSIFICATION 

        speech_type = classify_user_input(
            groq_client,
            user_text,
            current_q
        )

        # ==========================================================
        # BACKGROUND TALK
        # ==========================================================

        if speech_type == "BACKGROUND_TALK":
            if idx < len(q_bank):
                pending_q = q_bank[idx]
                st.session_state.current_question = pending_q
        
                reply = f"Please continue the interview. {pending_q}"
            else:
                reply = "Please continue the interview."

            st.session_state.chat_history.append({
                "role": "assistant",
                "content": reply
            })

            st.session_state.pending_voice = reply
            st.rerun()

        # ==========================================================
        # INTERRUPTION
        # ==========================================================

        if speech_type == "INTERRUPTION":
            reply = f"Sure. Let me repeat.\n\n{current_q}"

            st.session_state.chat_history.append({
                "role": "assistant",
                "content": reply
            })

            st.session_state.pending_voice = reply
            st.rerun()

        # ==========================================================
        # NEXT QUESTION
        # ==========================================================
        
        if speech_type == "NEXT_QUESTION":
            
            if idx < len(q_bank):
                
                next_q = q_bank[idx]

                st.session_state.current_question = next_q
                
                st.session_state.q_index += 1
        
                reply = f"Sure. Skipping this one. Here is your next question: {next_q}"

            else:
                reply = "Interview completed. Generating your report..."
                st.session_state.report_ready = True
        
            st.session_state.chat_history.append({
                "role": "assistant",
                "content": reply
            })
        
            st.session_state.pending_voice = reply
            st.rerun()

        # ==========================================================
        # TECHNICAL FLOW ONLY (SAFE)
        # ==========================================================
        
        existing = None

        if idx >= len(q_bank):
            reply = "Interview is already complete. Generating your report..."
            st.session_state.report_ready = True
        
        else:
            clean_answer = user_text.strip()

            if clean_answer not in ["...", ".", ""]:  
                existing = next((i for i, a in enumerate(st.session_state.answers) if a["question"] == current_q), None)

            if existing is not None:
                st.session_state.answers[existing]["followups"].append(user_text)
            else:
                st.session_state.answers.append({
                    "question": current_q,
                    "main_answer": user_text,
                    "followups": [],
                    "final_score": 0.0
                })

            next_q = q_bank[idx] if idx < len(q_bank) else "End"

            ai_reply = get_ai_decision(
                groq_client,
                user_text,
                next_q,
                st.session_state.chat_history
            )

            # Check if moving forward
            if st.session_state.followup_count >= 1 or "next question" in ai_reply.lower():    
                st.session_state.q_index += 1
                st.session_state.followup_count = 0
                
                if st.session_state.q_index < len(q_bank):
                    st.session_state.current_question = q_bank[st.session_state.q_index]
            else:
                st.session_state.followup_count += 1


            if st.session_state.followup_count >= 2:
                ai_reply = f"I see. Let's move to the next one to stay on track. {next_q}"
                st.session_state.q_index += 1
                st.session_state.followup_count = 0

                if st.session_state.q_index < len(q_bank):
                    st.session_state.current_question = q_bank[st.session_state.q_index]

            st.session_state.chat_history.append({
                "role": "assistant",
                "content": ai_reply
            })

            st.session_state.pending_voice = ai_reply
            st.rerun()


# ================== MAIN APP ==================
def main():
    st.set_page_config(
        page_title="🎙️ AI-Powered Voice-Based Interview Assistant",
        page_icon="🎤",
        layout="centered"
    )
    
    st.markdown("""
    <style>
        .main {
            background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        }
        .stButton > button {
            border-radius: 10px;
            font-weight: 600;
        }
    </style>
    """, unsafe_allow_html=True)
    
    st.title("🎙️ AI Voice Interview")
    
    # Initialize
    init_session()
    stt_model, groq_client = load_agent_engines()
    
    # ===== UPLOAD SCREEN =====
    if not st.session_state.is_started:
        render_upload_screen(groq_client)
        return
    

    if st.session_state.pending_voice:
        st.session_state.is_ai_speaking = True

        ai_voice_output(st.session_state.pending_voice)
        wait_time = max(3, len(st.session_state.pending_voice) / 8)

        st.session_state.ai_end_time = time.time() + wait_time

        # time.sleep(wait_time)

        st.session_state.pending_voice = None

        st.session_state.mic_counter += 1 

        # st.rerun()
    
    # Control panel
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        progress = min(st.session_state.q_index, len(st.session_state.q_bank))
        total = len(st.session_state.q_bank)
        st.metric("Progress", f"{progress}/{total}")
    
    with col3:
        if st.button("🛑 End Interview", type="secondary"):
            st.session_state.force_end = True
            st.rerun()
    
    if st.session_state.force_end:
        
        st.session_state.chat_history.append({
            "role": "assistant",
            "content": "Interview ended early. Generating report..."
        })
        pdf = generate_report()
        st.session_state.report_ready = True
        st.session_state.report_file = pdf
        st.session_state.pending_voice = "Interview ended. Your report is ready."
        st.session_state.force_end = False
        st.rerun()
        
    st.markdown("---")

    # MULTIPLE VOICE CHECK
    multiple_voice = st.query_params.get("multiple_voice", "false")
    
    if multiple_voice == "true":
    
        st.session_state.multiple_voice_count += 1
    
        st.session_state.cheating_flags.append(
            "Possible multiple voices detected during interview"
        )
    
        st.warning("⚠️ Multiple voices detected. Please ensure only candidate speaks.")
    
        st.query_params.clear()
    
    
    # TAB SWITCH DETECTION
    switched = st.query_params.get("tab_switch", "false")
    
    if switched == "true":
    
        st.session_state.tab_switch_count += 1
    
        st.session_state.suspicion_score += 1
    
        st.session_state.cheating_flags.append(
            "Candidate switched browser tab during interview"
        )
    
        st.warning("⚠️ Tab switching detected.")
    
        st.query_params.clear()
    
    # AI speech finished check

    if (
        st.session_state.get("is_ai_speaking")
        and time.time() > st.session_state.get("ai_end_time", 0)
    ):
        st.session_state.is_ai_speaking = False
    
    if not st.session_state.report_ready: 
        if st.session_state.pending_voice is None:
            st.write("### 🎤 AI is listening... (Speak now)")
           
            audio_bytes = audio_recorder(
                text="Listening...",
                recording_color="#e74c3c",
                neutral_color="#3498db",
                icon_name="microphone",
                icon_size="2x",
                pause_threshold=2.5, 
                sample_rate=16000,
                auto_start=True,
                key=f"mic_{st.session_state.mic_counter}" 
            )

        if audio_bytes:
            process_user_audio(audio_bytes, stt_model, groq_client)
        
    # ===== DOWNLOAD REPORT =====
    if st.session_state.report_ready:
        
        if not st.session_state.report_file:
            with st.spinner("Calculating final scores and generating report..."):
                st.session_state.report_file = generate_report()
        
        st.success("✅ Interview Completed!")
        
        total = sum(st.session_state.scores)
        maximum = len(st.session_state.q_bank) * 5
        percentage = (total / maximum * 100) if maximum > 0 else 0
        
        st.metric("Final Score", f"{total}/{maximum}", f"{percentage:.1f}%")
        
        with open(st.session_state.report_file, "rb") as f:
            st.download_button(
                label="📄 Download Report ",
                data=f,
                file_name="interview_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )


if __name__ == "__main__":
    main()
